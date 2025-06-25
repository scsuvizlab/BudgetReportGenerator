"""
Document Generator - FIXED VERSION

Fixed regex escaping issue in placeholder replacement.
"""
import re
from pathlib import Path
from typing import Dict, Optional, Any
from docx import Document
from docx.shared import Inches
import logging
from datetime import datetime

from template_document import TemplateDocument, Placeholder
from session_state import SessionState, FieldMapping

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Generates final budget justification documents."""
    
    def __init__(self):
        self.last_generated_path: Optional[Path] = None
    
    def generate_document(self, session: SessionState, output_path: Optional[Path] = None) -> Path:
        """Generate a document from the current session state."""
        if not session.template:
            raise ValueError("No template loaded")
        
        if not session.is_ready_for_generation():
            logger.warning("Session may not be ready for generation - proceeding anyway")
        
        # Determine output path
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"budget_justification_{timestamp}.{session.config.output_format}"
            output_path = session.config.output_directory / filename
        
        # Generate based on format
        if session.config.output_format.lower() == "docx":
            return self._generate_docx(session, output_path)
        elif session.config.output_format.lower() == "md":
            return self._generate_markdown(session, output_path)
        else:
            raise ValueError(f"Unsupported output format: {session.config.output_format}")
    
    def _generate_docx(self, session: SessionState, output_path: Path) -> Path:
        """Generate a Word document."""
        try:
            # Create or copy template document
            if session.template.source_type == "docx":
                # Copy the original DOCX and modify it
                doc = Document(session.template.source_path)
            else:
                # Create new DOCX from text content
                doc = Document()
                
                # Add title
                title = doc.add_heading('Budget Justification', 0)
                
                # Add content paragraphs
                content = self._replace_placeholders(session.template.content, session.mappings)
                for paragraph_text in content.split('\n'):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text)
            
            # Replace placeholders in existing DOCX
            if session.template.source_type == "docx":
                self._replace_placeholders_in_docx(doc, session.mappings)
            
            # Add generation metadata
            self._add_metadata_section(doc, session)
            
            # Save document
            doc.save(output_path)
            self.last_generated_path = output_path
            
            logger.info(f"DOCX document generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            raise
    
    def _generate_markdown(self, session: SessionState, output_path: Path) -> Path:
        """Generate a Markdown document."""
        try:
            # Replace placeholders in template content
            content = self._replace_placeholders(session.template.content, session.mappings)
            
            # Add metadata header
            metadata = self._create_metadata_section(session)
            
            # Combine content
            full_content = f"""# Budget Justification

{content}

---

{metadata}
"""
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_content)
            
            self.last_generated_path = output_path
            logger.info(f"Markdown document generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate Markdown: {e}")
            raise
    
    def _replace_placeholders(self, content: str, mappings: Dict[str, FieldMapping]) -> str:
        """Replace placeholders in text content with actual values - FIXED VERSION."""
        result = content
        
        for field_name, mapping in mappings.items():
            value = mapping.final_value
            
            if value is not None:
                # Format the value
                formatted_value = self._format_value(value)
                
                # FIXED: Properly escape field names for regex
                try:
                    # Handle different placeholder formats with proper escaping
                    if field_name.startswith('{') and field_name.endswith('}'):
                        # Field name already includes braces like {Project_Title}
                        # Use re.escape to properly escape all special regex characters
                        escaped_field = re.escape(field_name)
                        result = re.sub(escaped_field, formatted_value, result)
                    else:
                        # Field name without braces, try multiple patterns
                        # Use re.escape on the field name portion
                        escaped_field = re.escape(field_name)
                        
                        patterns = [
                            f"\\[{escaped_field}\\]",           # [PI_SALARY]
                            f"\\{{{escaped_field}\\}}",         # {PI_SALARY}  
                            f"\\$\\{{{escaped_field}\\}}",      # ${PI_SALARY}
                        ]
                        
                        for pattern in patterns:
                            result = re.sub(pattern, formatted_value, result, flags=re.IGNORECASE)
                    
                    # Handle implicit placeholders based on the placeholder's position
                    if hasattr(mapping, 'placeholder') and mapping.placeholder.placeholder_type == "implicit":
                        # Find the position and replace the pattern
                        start_pos = mapping.placeholder.position
                        context = mapping.placeholder.context
                        
                        # Simple approach: replace common patterns near this position
                        implicit_patterns = [
                            r'XXX\.XXX',
                            r'XXX,XXX',
                            r'____+',
                            r'\$\s*_+'
                        ]
                        
                        for pattern in implicit_patterns:
                            # Replace first occurrence in the context area
                            result = re.sub(pattern, formatted_value, result, count=1)
                
                except re.error as regex_err:
                    # If regex fails, try simple string replacement as fallback
                    logger.warning(f"Regex replacement failed for field '{field_name}': {regex_err}. Using simple replacement.")
                    result = result.replace(field_name, formatted_value)
                    
            else:
                # No value available - leave placeholder or mark as missing
                missing_text = "[VALUE NOT FOUND]"
                
                try:
                    if field_name.startswith('{') and field_name.endswith('}'):
                        # Field name already includes braces
                        escaped_field = re.escape(field_name)
                        result = re.sub(escaped_field, missing_text, result)
                    else:
                        escaped_field = re.escape(field_name)
                        patterns = [
                            f"\\[{escaped_field}\\]",
                            f"\\{{{escaped_field}\\}}",
                            f"\\$\\{{{escaped_field}\\}}"
                        ]
                        
                        for pattern in patterns:
                            result = re.sub(pattern, missing_text, result, flags=re.IGNORECASE)
                            
                except re.error as regex_err:
                    # Fallback to simple string replacement
                    logger.warning(f"Regex replacement failed for missing field '{field_name}': {regex_err}. Using simple replacement.")
                    result = result.replace(field_name, missing_text)
        
        return result
    
    def _replace_placeholders_in_docx(self, doc: Document, mappings: Dict[str, FieldMapping]) -> None:
        """Replace placeholders in a DOCX document."""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                original_text = paragraph.text
                new_text = self._replace_placeholders(original_text, mappings)
                
                if new_text != original_text:
                    # Replace the entire paragraph text
                    paragraph.clear()
                    paragraph.add_run(new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text:
                            original_text = paragraph.text
                            new_text = self._replace_placeholders(original_text, mappings)
                            
                            if new_text != original_text:
                                paragraph.clear()
                                paragraph.add_run(new_text)
    
    def _format_value(self, value: float) -> str:
        """Format a numeric value for display."""
        # Format as currency with commas
        if value >= 1000:
            return f"${value:,.0f}"
        else:
            return f"${value:.2f}"
    
    def _add_metadata_section(self, doc: Document, session: SessionState) -> None:
        """Add metadata section to Word document."""
        # Add page break
        doc.add_page_break()
        
        # Add metadata heading
        doc.add_heading('Generation Metadata', level=1)
        
        # Add metadata content
        metadata_text = self._create_metadata_section(session)
        doc.add_paragraph(metadata_text)
    
    def _create_metadata_section(self, session: SessionState) -> str:
        """Create metadata section text."""
        summary = session.get_mapping_summary()
        
        metadata = f"""
## Document Generation Information

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**Session ID:** {session.session_id}

### Template Information
- **Source:** {session.template.source_path.name if session.template else 'Unknown'}
- **Type:** {session.template.source_type if session.template else 'Unknown'}
- **Placeholders:** {len(session.template.placeholders) if session.template else 0}

### Budget Information
- **Source:** {session.budget.source_path.name if session.budget else 'Unknown'}
- **Sheets:** {len(session.budget.sheets) if session.budget else 0}
- **Budget Cells:** {len(session.budget.cells) if session.budget else 0}

### Field Mapping Summary
- **Total Fields:** {summary['total_fields']}
- **Mapped Fields:** {summary['mapped_fields']}
- **Manual Overrides:** {summary['manual_overrides']}
- **High Confidence:** {summary['high_confidence']}
- **Low Confidence:** {summary['low_confidence']}
- **Unmapped:** {summary['unmapped']}

### Usage Statistics
- **Total Tokens:** {session.total_tokens_used:,}
- **Total Cost:** ${session.total_cost_usd:.4f}

### Field Details
"""
        
        # Add field-by-field details
        for field_name, mapping in session.mappings.items():
            status = "✓" if mapping.final_value is not None else "✗"
            value_str = mapping.display_value
            source = "Manual" if mapping.is_manually_set else "Auto"
            confidence = f"{mapping.confidence:.2f}"
            
            metadata += f"- **{field_name}** {status} {value_str} ({source}, confidence: {confidence})\n"
            
            if mapping.notes:
                metadata += f"  - Notes: {mapping.notes}\n"
        
        return metadata
    
    def preview_content(self, session: SessionState) -> str:
        """Generate preview content without saving to file."""
        if not session.template:
            return "No template loaded"
        
        try:
            # Generate preview text
            content = self._replace_placeholders(session.template.content, session.mappings)
            
            # Add preview header
            preview = f"""
# Budget Justification Preview
*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*

---

{content}

---

*Preview generated from {session.template.source_path.name}*
"""
            return preview
            
        except Exception as e:
            logger.error(f"Failed to generate preview: {e}")
            return f"Error generating preview: {e}"
    
    def get_replacement_summary(self, session: SessionState) -> Dict[str, Any]:
        """Get a summary of what replacements will be made."""
        if not session.template:
            return {}
        
        summary = {
            "total_placeholders": len(session.template.placeholders),
            "successful_replacements": 0,
            "failed_replacements": 0,
            "manual_values": 0,
            "auto_values": 0,
            "replacements": []
        }
        
        for field_name, mapping in session.mappings.items():
            replacement_info = {
                "field": field_name,
                "has_value": mapping.final_value is not None,
                "value": mapping.display_value,
                "source": "manual" if mapping.is_manually_set else "auto",
                "confidence": mapping.confidence
            }
            
            summary["replacements"].append(replacement_info)
            
            if mapping.final_value is not None:
                summary["successful_replacements"] += 1
                
                if mapping.is_manually_set:
                    summary["manual_values"] += 1
                else:
                    summary["auto_values"] += 1
            else:
                summary["failed_replacements"] += 1
        
        return summary


# Testing function to validate regex patterns
def test_field_name_escaping():
    """Test function to validate the regex escaping works correctly."""
    import re
    
    test_cases = [
        "{PI_Salary}",
        "{Supplies_ADP/CS_PlasticSCM_Total}",
        "{Faculty_Train_the_trainer_1}",
        "[PI_SALARY]",
        "${TOTAL}",
        "test-range",
        "test[bracket]",
        "test(paren)",
        "test.dot",
        "test+plus"
    ]
    
    print("Testing regex escaping:")
    for field_name in test_cases:
        try:
            escaped = re.escape(field_name)
            pattern = f"\\{{{escaped}\\}}"
            test_text = f"Sample text with {field_name} placeholder"
            result = re.sub(pattern, "REPLACED", test_text)
            print(f"✓ {field_name} -> {escaped} (OK)")
        except re.error as e:
            print(f"✗ {field_name} -> Error: {e}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    # Run the test
    test_field_name_escaping()
    
    print("\nDocument Generator (Fixed Version) ready.")