"""
Template Document Parser

Handles parsing of various template formats (DOCX, MD, TXT, PDF)
and extracts placeholders that need to be filled with budget data.
"""
import re
import logging
from pathlib import Path
from typing import List, Dict, Optional, Any, Set
from dataclasses import dataclass
from docx import Document
import markdown

logger = logging.getLogger(__name__)


@dataclass
class Placeholder:
    """Represents a placeholder found in a template."""
    text: str  # The placeholder text including braces/brackets
    position: int  # Character position in document
    placeholder_type: str  # 'explicit' or 'implicit'
    context: str  # Surrounding text for context
    pattern: str = ""  # Original pattern that matched
    
    @property
    def clean_text(self) -> str:
        """Get placeholder text without braces/brackets."""
        text = self.text
        # Remove common placeholder markers
        for marker in ['{', '}', '[', ']', '$']:
            text = text.replace(marker, '')
        return text.strip()


@dataclass
class TemplateDocument:
    """Canonical representation of a template document."""
    source_path: Path
    source_type: str  # 'docx', 'md', 'txt', 'pdf'
    content: str  # Full text content
    placeholders: List[Placeholder]
    tables: List[Dict[str, Any]] = None  # For future table support
    position_map: Dict[str, int] = None  # Maps placeholder to positions
    
    def get_placeholders_by_type(self, placeholder_type: str) -> List[Placeholder]:
        """Get placeholders of a specific type."""
        return [p for p in self.placeholders if p.placeholder_type == placeholder_type]
    
    def find_placeholder(self, text: str) -> Optional[Placeholder]:
        """Find a placeholder by its text."""
        for placeholder in self.placeholders:
            if placeholder.text == text or placeholder.clean_text == text:
                return placeholder
        return None


class TemplateParser:
    """Parses template documents and extracts placeholders."""
    
    # Common placeholder patterns
    EXPLICIT_PATTERNS = [
        r'\{[^}]+\}',           # {placeholder}
        r'\[[^\]]+\]',          # [placeholder]  
        r'\$\{[^}]+\}',         # ${placeholder}
        r'\{\{[^}]+\}\}',       # {{placeholder}}
        r'\<[^>]+\>',           # <placeholder>
    ]
    
    # Implicit placeholder patterns (things that look like they need values)
    IMPLICIT_PATTERNS = [
        r'\$\s*[_]{3,}',        # $_____ 
        r'[_]{5,}',             # _____
        r'XXX[,.]?XXX',         # XXX,XXX or XXX.XXX
        r'\$\s*\d+[,.]?\d*\s*\w*',  # $1,234 (example values)
        r'TBD|TBA|TODO',        # To be determined/announced/done
    ]
    
    def __init__(self):
        """Initialize the template parser."""
        self.context_window = 50  # Characters around placeholder for context
    
    def parse_file(self, file_path: Path) -> TemplateDocument:
        """Parse a template file and return TemplateDocument."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Template file not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return self._parse_docx(file_path)
        elif suffix == '.md':
            return self._parse_markdown(file_path)
        elif suffix in ['.txt', '.text']:
            return self._parse_text(file_path)
        elif suffix == '.pdf':
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"Unsupported template format: {suffix}")
    
    def _parse_docx(self, file_path: Path) -> TemplateDocument:
        """Parse a DOCX template."""
        try:
            doc = Document(file_path)
            
            # Extract text content
            content_parts = []
            for paragraph in doc.paragraphs:
                content_parts.append(paragraph.text)
            
            # Add table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content_parts.append(cell.text)
            
            content = '\n'.join(content_parts)
            
            # Find placeholders
            placeholders = self._extract_placeholders(content)
            
            logger.info(f"Parsed DOCX template: {len(placeholders)} placeholders found")
            
            return TemplateDocument(
                source_path=file_path,
                source_type='docx',
                content=content,
                placeholders=placeholders
            )
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX template {file_path}: {e}")
            raise
    
    def _parse_markdown(self, file_path: Path) -> TemplateDocument:
        """Parse a Markdown template."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Find placeholders in raw markdown
            placeholders = self._extract_placeholders(content)
            
            logger.info(f"Parsed Markdown template: {len(placeholders)} placeholders found")
            
            return TemplateDocument(
                source_path=file_path,
                source_type='md',
                content=content,
                placeholders=placeholders
            )
            
        except Exception as e:
            logger.error(f"Failed to parse Markdown template {file_path}: {e}")
            raise
    
    def _parse_text(self, file_path: Path) -> TemplateDocument:
        """Parse a plain text template."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            placeholders = self._extract_placeholders(content)
            
            logger.info(f"Parsed text template: {len(placeholders)} placeholders found")
            
            return TemplateDocument(
                source_path=file_path,
                source_type='txt',
                content=content,
                placeholders=placeholders
            )
            
        except Exception as e:
            logger.error(f"Failed to parse text template {file_path}: {e}")
            raise
    
    def _parse_pdf(self, file_path: Path) -> TemplateDocument:
        """Parse a PDF template (requires OCR or PDF text extraction)."""
        try:
            # Try to extract text from PDF
            content = self._extract_pdf_text(file_path)
            
            if not content.strip():
                raise ValueError("No text could be extracted from PDF")
            
            placeholders = self._extract_placeholders(content)
            
            logger.info(f"Parsed PDF template: {len(placeholders)} placeholders found")
            
            return TemplateDocument(
                source_path=file_path,
                source_type='pdf',
                content=content,
                placeholders=placeholders
            )
            
        except Exception as e:
            logger.error(f"Failed to parse PDF template {file_path}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            # Try using PyPDF2 first
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                
                return '\n'.join(text_parts)
                
        except ImportError:
            logger.warning("PyPDF2 not available, trying alternative PDF extraction")
            
        try:
            # Try using pdfplumber as alternative
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return '\n'.join(text_parts)
                
        except ImportError:
            logger.warning("pdfplumber not available either")
        
        # Fallback: return error message for manual handling
        return f"[PDF_CONTENT_PLACEHOLDER - Manual extraction required for {file_path.name}]"
    
    def _extract_placeholders(self, content: str) -> List[Placeholder]:
        """Extract all placeholders from content."""
        placeholders = []
        found_positions = set()  # Track positions to avoid duplicates
        
        # Find explicit placeholders
        for pattern in self.EXPLICIT_PATTERNS:
            for match in re.finditer(pattern, content, re.IGNORECASE):
                position = match.start()
                
                if position not in found_positions:
                    placeholder = Placeholder(
                        text=match.group(),
                        position=position,
                        placeholder_type='explicit',
                        context=self._get_context(content, position),
                        pattern=pattern
                    )
                    placeholders.append(placeholder)
                    found_positions.add(position)
        
        # Find implicit placeholders (but be more conservative)
        for pattern in self.IMPLICIT_PATTERNS:
            for match in re.finditer(pattern, content, re.IGNORECASE):
                position = match.start()
                
                if position not in found_positions:
                    # Additional validation for implicit placeholders
                    if self._is_likely_placeholder(content, match):
                        placeholder = Placeholder(
                            text=match.group(),
                            position=position,
                            placeholder_type='implicit',
                            context=self._get_context(content, position),
                            pattern=pattern
                        )
                        placeholders.append(placeholder)
                        found_positions.add(position)
        
        # Sort by position
        placeholders.sort(key=lambda p: p.position)
        
        # Remove duplicates and clean up
        placeholders = self._deduplicate_placeholders(placeholders)
        
        logger.info(f"Found {len(placeholders)} unique placeholders")
        return placeholders
    
    def _get_context(self, content: str, position: int) -> str:
        """Get surrounding context for a placeholder."""
        start = max(0, position - self.context_window)
        end = min(len(content), position + self.context_window)
        return content[start:end]
    
    def _is_likely_placeholder(self, content: str, match: re.Match) -> bool:
        """Check if an implicit match is likely a real placeholder."""
        text = match.group()
        context = self._get_context(content, match.start())
        
        # Skip if it looks like formatting or example text
        skip_patterns = [
            r'example',
            r'sample',
            r'test',
            r'lorem ipsum',
            r'draft',
            r'version',
            r'page \d+',
            r'figure \d+',
            r'table \d+'
        ]
        
        context_lower = context.lower()
        for skip_pattern in skip_patterns:
            if re.search(skip_pattern, context_lower):
                return False
        
        # Look for budget-related context
        budget_indicators = [
            r'cost',
            r'budget',
            r'salary',
            r'total',
            r'amount',
            r'price',
            r'expense',
            r'funding',
            r'grant',
            r'award'
        ]
        
        for indicator in budget_indicators:
            if re.search(indicator, context_lower):
                return True
        
        # Default: include if it's in a line that looks like it needs a value
        line_start = content.rfind('\n', 0, match.start()) + 1
        line_end = content.find('\n', match.end())
        if line_end == -1:
            line_end = len(content)
        
        line = content[line_start:line_end]
        
        # Check if line has budget-related keywords
        return any(re.search(indicator, line, re.IGNORECASE) for indicator in budget_indicators)
    
    def _deduplicate_placeholders(self, placeholders: List[Placeholder]) -> List[Placeholder]:
        """Remove duplicate placeholders and merge similar ones."""
        if not placeholders:
            return []
        
        # Group by clean text
        groups = {}
        for placeholder in placeholders:
            key = placeholder.clean_text.lower()
            if key not in groups:
                groups[key] = []
            groups[key].append(placeholder)
        
        # Keep the best placeholder from each group
        result = []
        for group in groups.values():
            if len(group) == 1:
                result.append(group[0])
            else:
                # Prefer explicit over implicit
                explicit = [p for p in group if p.placeholder_type == 'explicit']
                if explicit:
                    result.append(explicit[0])  # Take first explicit
                else:
                    result.append(group[0])  # Take first implicit
        
        # Sort by position again
        result.sort(key=lambda p: p.position)
        
        return result
    
    def preview_placeholders(self, template_doc: TemplateDocument) -> str:
        """Generate a preview of found placeholders."""
        preview_lines = [
            f"Template: {template_doc.source_path.name}",
            f"Type: {template_doc.source_type.upper()}",
            f"Placeholders found: {len(template_doc.placeholders)}",
            "",
            "Placeholders:"
        ]
        
        for i, placeholder in enumerate(template_doc.placeholders, 1):
            context_preview = placeholder.context.replace('\n', ' ').strip()
            if len(context_preview) > 80:
                context_preview = context_preview[:77] + "..."
            
            preview_lines.append(
                f"{i:2d}. {placeholder.text} ({placeholder.placeholder_type})"
            )
            preview_lines.append(f"    Context: {context_preview}")
        
        return '\n'.join(preview_lines)


# Example usage and testing
if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    
    parser = TemplateParser()
    
    # Test with the provided template
    test_content = """
    ## Project Title:
    {Project_Title}
    
    ## Principal Investigator:
    {Principal_Investigator}
    
    ### Personnel
    
    #### Principal Investigator:
    Name: {Principal_Investigator}
    Total Cost: {PI_IFO_1_Total} USD
    
    #### Co-Principal Investigator:
    Name: {Co_Principal_Investigator_Name_1}
    Total Cost: {CO_PI_IFO_1_Total} USD
    
    ### Travel
    Domestic Travel Year 1: {Travel_Domestic_Year_1} USD
    Total: {Travel_Domestic_Total} USD
    """
    
    # Create a mock template document
    from pathlib import Path
    import tempfile
    
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(test_content)
        temp_path = Path(f.name)
    
    try:
        template_doc = parser.parse_file(temp_path)
        
        print("Parsed template successfully!")
        print(f"Found {len(template_doc.placeholders)} placeholders:")
        
        for placeholder in template_doc.placeholders:
            print(f"  - {placeholder.text} ({placeholder.placeholder_type})")
        
        print("\nPreview:")
        print(parser.preview_placeholders(template_doc))
        
    finally:
        temp_path.unlink()  # Clean up temp file