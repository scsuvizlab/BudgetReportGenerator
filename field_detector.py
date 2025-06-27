"""
Enhanced field detector for template analysis with improved pattern recognition
for grant proposals and budget justifications. Better handles string-based fields
like names, titles, descriptions, and notes.
"""
import re
import logging
from typing import Dict, List, Tuple, Set, Optional, Any
from dataclasses import dataclass, field
from pathlib import Path
import docx
from docx.document import Document
try:
    import pypdf
    from pypdf import PdfReader
except ImportError:
    try:
        import PyPDF2 as pypdf
        from PyPDF2 import PdfReader
    except ImportError:
        pypdf = None
        PdfReader = None
import markdown
from rapidfuzz import fuzz, process

@dataclass
class TemplateField:
    """Represents a field found in a template document."""
    name: str
    placeholder_text: str
    field_type: str  # 'text', 'numeric', 'currency', 'personnel', 'description', 'date'
    context: str  # Surrounding text for context
    position: int  # Position in document
    confidence: float = 1.0
    variations: List[str] = field(default_factory=list)
    requirements: Dict[str, Any] = field(default_factory=dict)
    
    # Backward compatibility attributes for GUI
    source_type: str = "template"  # GUI expects this
    source: str = ""  # GUI might expect this
    description: str = ""  # GUI might expect this
    
    def __post_init__(self):
        """Set backward compatibility attributes after initialization."""
        if not self.source:
            self.source = self.placeholder_text
        if not self.description:
            self.description = self.context[:100] + "..." if len(self.context) > 100 else self.context

@dataclass
class FieldPattern:
    """Pattern information for field recognition."""
    pattern_type: str
    keywords: List[str]
    context_clues: List[str]
    field_type: str
    priority: int = 5  # 1-10, higher is more important

class FieldDetector:
    """Enhanced field detector with grant-specific pattern recognition."""
    
    def __init__(self, enable_debug=False):
        self.enable_debug = enable_debug
        self.logger = logging.getLogger(__name__)
        
        # Define field patterns for grant proposals
        self.field_patterns = [
            # Personnel fields
            FieldPattern(
                pattern_type='personnel_name',
                keywords=['investigator', 'pi', 'co-pi', 'name', 'researcher', 'scientist'],
                context_clues=['principal', 'co-investigator', 'staff', 'personnel', 'team member'],
                field_type='personnel',
                priority=9
            ),
            FieldPattern(
                pattern_type='personnel_title',
                keywords=['title', 'position', 'role', 'rank'],
                context_clues=['professor', 'assistant', 'associate', 'director', 'postdoc'],
                field_type='personnel',
                priority=8
            ),
            FieldPattern(
                pattern_type='personnel_effort',
                keywords=['effort', 'fte', 'time', 'percent', '%'],
                context_clues=['commitment', 'allocation', 'dedication'],
                field_type='numeric',
                priority=7
            ),
            
            # Budget and cost fields
            FieldPattern(
                pattern_type='salary_cost',
                keywords=['salary', 'wage', 'compensation', 'cost', 'amount'],
                context_clues=['annual', 'monthly', 'hourly', 'total', 'base'],
                field_type='currency',
                priority=9
            ),
            FieldPattern(
                pattern_type='equipment_cost',
                keywords=['equipment', 'instrument', 'computer', 'software', 'hardware'],
                context_clues=['purchase', 'cost', 'price', 'total'],
                field_type='currency',
                priority=8
            ),
            FieldPattern(
                pattern_type='travel_cost',
                keywords=['travel', 'conference', 'trip', 'transportation'],
                context_clues=['airfare', 'lodging', 'per diem', 'registration'],
                field_type='currency',
                priority=7
            ),
            
            # Description and justification fields
            FieldPattern(
                pattern_type='description',
                keywords=['description', 'explain', 'describe', 'detail'],
                context_clues=['purpose', 'reason', 'justification', 'rationale'],
                field_type='description',
                priority=8
            ),
            FieldPattern(
                pattern_type='notes',
                keywords=['notes', 'comments', 'remarks', 'additional'],
                context_clues=['information', 'details', 'explanation'],
                field_type='description',
                priority=6
            ),
            FieldPattern(
                pattern_type='justification',
                keywords=['justification', 'justify', 'rationale', 'reason'],
                context_clues=['necessary', 'required', 'essential', 'purpose'],
                field_type='description',
                priority=9
            ),
            
            # Project information
            FieldPattern(
                pattern_type='project_title',
                keywords=['project', 'title', 'study', 'research'],
                context_clues=['name', 'called', 'entitled'],
                field_type='text',
                priority=9
            ),
            FieldPattern(
                pattern_type='institution',
                keywords=['institution', 'university', 'organization', 'company'],
                context_clues=['affiliation', 'employer', 'workplace'],
                field_type='text',
                priority=7
            ),
            
            # Time-related fields
            FieldPattern(
                pattern_type='duration',
                keywords=['duration', 'period', 'timeline', 'months', 'years'],
                context_clues=['project', 'study', 'research', 'length'],
                field_type='numeric',
                priority=6
            ),
            FieldPattern(
                pattern_type='dates',
                keywords=['date', 'start', 'end', 'begin', 'finish'],
                context_clues=['project', 'period', 'timeline'],
                field_type='date',
                priority=6
            )
        ]
        
        # Common placeholder patterns
        self.placeholder_patterns = [
            r'\{([^}]+)\}',  # {field_name}
            r'\[\[([^\]]+)\]\]',  # [[field_name]]
            r'\[([^\]]+)\]',  # [field_name]
            r'<<([^>]+)>>',  # <<field_name>>
            r'{{([^}]+)}}',  # {{field_name}}
            r'<([^>]+)>',  # <field_name>
            r'_+([A-Za-z_\s]+)_+',  # ___field_name___
            r'\.\.\.([A-Za-z_\s]+)\.\.\.',  # ...field_name...
        ]
        
        # Grant-specific terminology
        self.grant_terminology = {
            'personnel_categories': [
                'principal investigator', 'pi', 'co-principal investigator', 'co-pi',
                'co-investigator', 'senior personnel', 'key personnel', 'postdoc',
                'graduate student', 'undergraduate student', 'research scientist',
                'research associate', 'technician', 'staff scientist'
            ],
            'cost_categories': [
                'personnel', 'fringe benefits', 'equipment', 'travel', 'participant support',
                'other direct costs', 'total direct costs', 'facilities and administrative',
                'indirect costs', 'total project costs', 'supplies', 'contractual'
            ],
            'time_units': [
                'calendar months', 'academic months', 'summer months', 'person months',
                'full-time equivalent', 'fte', 'percent effort', '% effort'
            ]
        }
    
    def analyze_template(self, template_path: str) -> List[TemplateField]:
        """Analyze a template document to identify fields that need to be filled."""
        template_path = Path(template_path)
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        self.logger.info(f"Analyzing template: {template_path}")
        
        # Extract text based on file type
        if template_path.suffix.lower() == '.docx':
            text = self._extract_docx_text(template_path)
        elif template_path.suffix.lower() == '.pdf':
            text = self._extract_pdf_text(template_path)
        elif template_path.suffix.lower() in ['.md', '.markdown']:
            text = self._extract_markdown_text(template_path)
        elif template_path.suffix.lower() == '.txt':
            text = self._extract_txt_text(template_path)
        else:
            raise ValueError(f"Unsupported template format: {template_path.suffix}")
        
        # Find explicit placeholders
        explicit_fields = self._find_explicit_placeholders(text)
        
        # Find implicit fields through pattern analysis
        implicit_fields = self._find_implicit_fields(text)
        
        # Combine and deduplicate fields
        all_fields = self._merge_field_lists(explicit_fields, implicit_fields)
        
        # Enhance fields with context and type information
        enhanced_fields = self._enhance_fields(all_fields, text)
        
        # Sort by priority and confidence
        enhanced_fields.sort(key=lambda x: (x.confidence, -x.position), reverse=True)
        
        self.logger.info(f"Found {len(enhanced_fields)} template fields")
        return enhanced_fields
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from a Word document."""
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            return '\n'.join(paragraphs)
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from a PDF document."""
        if not pypdf or not PdfReader:
            raise ImportError("PDF support requires 'pypdf' or 'PyPDF2' package. Install with: pip install pypdf")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                text_parts = []
                
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                
                return '\n'.join(text_parts)
        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {e}")
            raise
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        """Extract text from a Markdown document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Convert markdown to plain text (basic approach)
            # Remove markdown syntax
            content = re.sub(r'^#+\s+', '', content, flags=re.MULTILINE)  # Headers
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Bold
            content = re.sub(r'\*(.*?)\*', r'\1', content)  # Italic
            content = re.sub(r'`(.*?)`', r'\1', content)  # Inline code
            content = re.sub(r'^\s*-\s+', '', content, flags=re.MULTILINE)  # List items
            
            return content
        except Exception as e:
            self.logger.error(f"Error extracting Markdown text: {e}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _find_explicit_placeholders(self, text: str) -> List[TemplateField]:
        """Find explicit placeholder patterns in the text."""
        fields = []
        position = 0
        
        for pattern in self.placeholder_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            
            for match in matches:
                field_name = match.group(1).strip()
                placeholder_text = match.group(0)
                
                # Get context around the placeholder
                start_pos = max(0, match.start() - 50)
                end_pos = min(len(text), match.end() + 50)
                context = text[start_pos:end_pos]
                
                # Determine field type based on name and context
                field_type = self._determine_field_type(field_name, context)
                
                field = TemplateField(
                    name=field_name,
                    placeholder_text=placeholder_text,
                    field_type=field_type,
                    context=context,
                    position=match.start(),
                    confidence=0.95
                )
                
                fields.append(field)
                position += 1
        
        return fields
    
    def _find_implicit_fields(self, text: str) -> List[TemplateField]:
        """Find implicit fields through pattern analysis and context."""
        fields = []
        sentences = re.split(r'[.!?]+', text)
        
        for i, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Look for patterns that suggest missing information
            implicit_patterns = [
                r'(name|title|position)\s*[:;]?\s*$',  # Ends with field label
                r'(amount|cost|total|salary)\s*[:;]?\s*\$?\s*$',  # Ends with money field
                r'(description|explain|describe|details?)\s*[:;]?\s*$',  # Ends with description request
                r'(note|comment|remark)s?\s*[:;]?\s*$',  # Ends with notes request
                r'(justify|justification|rationale)\s*[:;]?\s*$',  # Ends with justification request
            ]
            
            for pattern in implicit_patterns:
                match = re.search(pattern, sentence, re.IGNORECASE)
                if match:
                    field_name = match.group(1)
                    field_type = self._determine_field_type(field_name, sentence)
                    
                    # Calculate position in original text
                    position = text.find(sentence)
                    
                    field = TemplateField(
                        name=field_name.lower(),
                        placeholder_text=f"[{field_name}]",
                        field_type=field_type,
                        context=sentence,
                        position=position if position >= 0 else i * 100,
                        confidence=0.7
                    )
                    
                    fields.append(field)
        
        # Also look for grant-specific patterns
        grant_fields = self._find_grant_specific_patterns(text)
        fields.extend(grant_fields)
        
        return fields
    
    def _find_grant_specific_patterns(self, text: str) -> List[TemplateField]:
        """Find patterns specific to grant proposals and budget justifications."""
        fields = []
        
        # Look for personnel-related blanks or incomplete information
        personnel_patterns = [
            r'(principal investigator|pi)\s*[:;]?\s*([A-Z][a-z]+\s+[A-Z][a-z]+)?\s*$',
            r'(co-?investigator|co-?pi)\s*[:;]?\s*([A-Z][a-z]+\s+[A-Z][a-z]+)?\s*$',
            r'(postdoc|graduate student|staff)\s*[:;]?\s*([A-Z][a-z]+\s+[A-Z][a-z]+)?\s*$',
        ]
        
        for pattern in personnel_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                role = match.group(1)
                name = match.group(2) if len(match.groups()) > 1 else None
                
                if not name:  # Missing name - this is a field
                    field = TemplateField(
                        name=f"{role.lower()}_name",
                        placeholder_text=f"[{role} Name]",
                        field_type='personnel',
                        context=match.group(0),
                        position=match.start(),
                        confidence=0.8
                    )
                    fields.append(field)
        
        # Look for cost tables or budget sections with missing values
        budget_patterns = [
            r'(salary|wage|cost|amount|total)\s*[:;]?\s*\$?\s*_+',  # Salary: ____
            r'(equipment|supplies|travel)\s*[:;]?\s*\$?\s*_+',  # Equipment: ____
            r'(indirect|overhead)\s*rate\s*[:;]?\s*_+\s*%?',  # Indirect rate: ___%
        ]
        
        for pattern in budget_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                category = match.group(1)
                field_type = 'currency' if 'rate' not in match.group(0).lower() else 'numeric'
                
                field = TemplateField(
                    name=f"{category.lower()}_amount",
                    placeholder_text=f"[{category} Amount]",
                    field_type=field_type,
                    context=match.group(0),
                    position=match.start(),
                    confidence=0.75
                )
                fields.append(field)
        
        return fields
    
    def _determine_field_type(self, field_name: str, context: str) -> str:
        """Determine the type of field based on name and context."""
        field_name_lower = field_name.lower()
        context_lower = context.lower()
        
        # Check against known patterns
        for pattern in self.field_patterns:
            # Check if field name matches pattern keywords
            name_match = any(keyword in field_name_lower for keyword in pattern.keywords)
            
            # Check if context matches pattern context clues
            context_match = any(clue in context_lower for clue in pattern.context_clues)
            
            if name_match or context_match:
                return pattern.field_type
        
        # Fallback logic
        if any(term in field_name_lower for term in ['name', 'investigator', 'pi', 'staff']):
            return 'personnel'
        elif any(term in field_name_lower for term in ['cost', 'amount', 'salary', 'budget', '$']):
            return 'currency'
        elif any(term in field_name_lower for term in ['percent', '%', 'rate', 'effort', 'fte']):
            return 'numeric'
        elif any(term in field_name_lower for term in ['description', 'notes', 'explain', 'justify']):
            return 'description'
        elif any(term in field_name_lower for term in ['date', 'start', 'end', 'begin']):
            return 'date'
        else:
            return 'text'
    
    def _merge_field_lists(self, explicit_fields: List[TemplateField], 
                          implicit_fields: List[TemplateField]) -> List[TemplateField]:
        """Merge explicit and implicit field lists, removing duplicates."""
        all_fields = explicit_fields.copy()
        
        for implicit_field in implicit_fields:
            # Check if this field is already covered by an explicit field
            is_duplicate = False
            
            for explicit_field in explicit_fields:
                # Check name similarity
                name_similarity = fuzz.ratio(implicit_field.name.lower(), 
                                           explicit_field.name.lower())
                
                # Check position proximity (within 100 characters)
                position_close = abs(implicit_field.position - explicit_field.position) < 100
                
                if name_similarity > 80 or position_close:
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                all_fields.append(implicit_field)
        
        return all_fields
    
    def _enhance_fields(self, fields: List[TemplateField], text: str) -> List[TemplateField]:
        """Enhance fields with additional context and requirements."""
        enhanced_fields = []
        
        for field in fields:
            # Generate variations of the field name
            variations = self._generate_field_variations(field.name)
            field.variations = variations
            
            # Determine requirements based on field type and context
            requirements = self._determine_field_requirements(field, text)
            field.requirements = requirements
            
            # Adjust confidence based on context quality
            field.confidence = self._calculate_field_confidence(field, text)
            
            enhanced_fields.append(field)
        
        return enhanced_fields
    
    def _generate_field_variations(self, field_name: str) -> List[str]:
        """Generate variations of a field name for better matching."""
        variations = [field_name]
        
        # Convert between naming conventions
        # snake_case to camelCase
        camel_case = re.sub(r'_([a-z])', lambda m: m.group(1).upper(), field_name)
        if camel_case != field_name:
            variations.append(camel_case)
        
        # snake_case to Title Case
        title_case = field_name.replace('_', ' ').title()
        if title_case != field_name:
            variations.append(title_case)
        
        # Add common synonyms
        synonyms = {
            'name': ['full_name', 'person_name', 'investigator_name'],
            'title': ['position', 'role', 'job_title'],
            'cost': ['amount', 'price', 'total', 'expense'],
            'description': ['desc', 'details', 'explanation'],
            'notes': ['comments', 'remarks', 'additional_info'],
            'effort': ['fte', 'time_commitment', 'percent_effort'],
            'salary': ['wage', 'compensation', 'pay'],
        }
        
        for base_term, synonym_list in synonyms.items():
            if base_term in field_name.lower():
                for synonym in synonym_list:
                    variations.append(field_name.lower().replace(base_term, synonym))
        
        return list(set(variations))  # Remove duplicates
    
    def _determine_field_requirements(self, field: TemplateField, text: str) -> Dict[str, Any]:
        """Determine requirements and constraints for a field."""
        requirements = {
            'required': True,
            'data_type': field.field_type,
            'max_length': None,
            'format_pattern': None,
            'validation_rules': []
        }
        
        # Set requirements based on field type
        if field.field_type == 'personnel':
            requirements.update({
                'max_length': 100,
                'validation_rules': ['must_be_person_name'],
                'format_pattern': r'^[A-Za-z\s\.,\-]+$'
            })
        elif field.field_type == 'currency':
            requirements.update({
                'validation_rules': ['must_be_positive_number'],
                'format_pattern': r'^\$?[\d,]+\.?\d*$'
            })
        elif field.field_type == 'numeric':
            if 'percent' in field.name.lower() or '%' in field.context:
                requirements.update({
                    'validation_rules': ['must_be_percentage'],
                    'format_pattern': r'^\d+\.?\d*%?$'
                })
        elif field.field_type == 'description':
            requirements.update({
                'max_length': 1000,
                'validation_rules': ['must_not_be_empty']
            })
        elif field.field_type == 'date':
            requirements.update({
                'format_pattern': r'^\d{1,2}/\d{1,2}/\d{4}$',
                'validation_rules': ['must_be_valid_date']
            })
        
        # Check context for additional requirements
        context_lower = field.context.lower()
        if 'required' in context_lower or 'mandatory' in context_lower:
            requirements['required'] = True
        elif 'optional' in context_lower:
            requirements['required'] = False
        
        return requirements
    
    def _calculate_field_confidence(self, field: TemplateField, text: str) -> float:
        """Calculate confidence score for field detection."""
        confidence = field.confidence
        
        # Adjust based on field name clarity
        if len(field.name) < 3:
            confidence *= 0.8  # Very short names are less reliable
        elif len(field.name) > 50:
            confidence *= 0.9  # Very long names might be noise
        
        # Adjust based on context quality
        context_words = len(field.context.split())
        if context_words < 3:
            confidence *= 0.9  # Little context
        elif context_words > 20:
            confidence *= 1.1  # Rich context
        
        # Boost confidence for grant-specific terminology
        context_lower = field.context.lower()
        for category, terms in self.grant_terminology.items():
            if any(term in context_lower for term in terms):
                confidence *= 1.1
                break
        
        # Penalize fields that might be false positives
        false_positive_indicators = [
            'example', 'sample', 'template', 'placeholder', 'dummy'
        ]
        if any(indicator in field.name.lower() for indicator in false_positive_indicators):
            confidence *= 0.7
        
        return min(confidence, 1.0)
    
    def suggest_field_mappings(self, template_fields: List[TemplateField], 
                             budget_text_values: List[Tuple[str, str, Any]]) -> Dict[str, List[Tuple[str, float]]]:
        """Suggest mappings between template fields and budget values."""
        suggestions = {}
        
        for field in template_fields:
            field_suggestions = []
            
            # For each text value in the budget, calculate match score
            for location, sheet_name, value in budget_text_values:
                if not isinstance(value, str) or not value.strip():
                    continue
                
                match_score = self._calculate_field_match_score(field, value)
                
                if match_score > 0.3:  # Minimum threshold
                    field_suggestions.append((location, match_score))
            
            # Sort by match score
            field_suggestions.sort(key=lambda x: x[1], reverse=True)
            
            # Keep top 5 suggestions
            suggestions[field.name] = field_suggestions[:5]
        
        return suggestions
    
    def _calculate_field_match_score(self, field: TemplateField, value: str) -> float:
        """Calculate how well a budget value matches a template field."""
        score = 0.0
        
        # Direct name matching
        for field_variation in [field.name] + field.variations:
            similarity = fuzz.partial_ratio(field_variation.lower(), value.lower())
            score = max(score, similarity / 100.0 * 0.4)
        
        # Type-based matching
        if field.field_type == 'personnel':
            if self._looks_like_person_name(value):
                score += 0.6
        elif field.field_type == 'currency':
            if self._looks_like_currency(value):
                score += 0.6
        elif field.field_type == 'description':
            if len(value) > 50:  # Long text likely to be description
                score += 0.4
        elif field.field_type == 'numeric':
            if self._looks_like_number(value):
                score += 0.5
        
        # Context-based matching
        # Check if value contains words from field context
        field_context_words = set(re.findall(r'\w+', field.context.lower()))
        value_words = set(re.findall(r'\w+', value.lower()))
        
        common_words = field_context_words.intersection(value_words)
        if common_words:
            context_score = len(common_words) / max(len(field_context_words), 1)
            score += context_score * 0.3
        
        return min(score, 1.0)
    
    def _looks_like_person_name(self, value: str) -> bool:
        """Check if value looks like a person's name."""
        if not isinstance(value, str) or len(value.strip()) < 3:
            return False
        
        # Simple heuristics
        words = value.strip().split()
        if len(words) < 2 or len(words) > 4:
            return False
        
        # Names typically start with capital letters
        if not all(word[0].isupper() for word in words if word):
            return False
        
        # Names shouldn't contain numbers
        if any(char.isdigit() for char in value):
            return False
        
        return True
    
    def _looks_like_currency(self, value: str) -> bool:
        """Check if value looks like a currency amount."""
        currency_patterns = [
            r'\$[\d,]+\.?\d*',
            r'[\d,]+\.?\d*\s*\$',
            r'USD\s*[\d,]+\.?\d*',
            r'[\d,]+\.?\d*\s*dollars?'
        ]
        
        return any(re.search(pattern, value, re.IGNORECASE) for pattern in currency_patterns)
    
    def _looks_like_number(self, value: str) -> bool:
        """Check if value looks like a number."""
        try:
            # Remove common formatting
            cleaned = re.sub(r'[,%$]', '', value.strip())
            float(cleaned)
            return True
        except ValueError:
            return False
    
    def get_field_summary(self, fields: List[TemplateField]) -> Dict[str, Any]:
        """Get a summary of detected fields for reporting."""
        type_counts = {}
        confidence_levels = {'high': 0, 'medium': 0, 'low': 0}
        
        for field in fields:
            # Count by type
            field_type = field.field_type
            type_counts[field_type] = type_counts.get(field_type, 0) + 1
            
            # Count by confidence level
            if field.confidence >= 0.8:
                confidence_levels['high'] += 1
            elif field.confidence >= 0.5:
                confidence_levels['medium'] += 1
            else:
                confidence_levels['low'] += 1
        
        return {
            'total_fields': len(fields),
            'field_types': type_counts,
            'confidence_distribution': confidence_levels,
            'high_priority_fields': [f.name for f in fields if f.confidence >= 0.8],
            'requires_attention': [f.name for f in fields if f.confidence < 0.5]
        }